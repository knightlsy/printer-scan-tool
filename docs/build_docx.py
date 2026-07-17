# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\EDY\WorkBuddy\2026-07-07-10-57-44\printer-scan-tool\docs"
IMG = os.path.join(BASE, "images")
OUT = os.path.join(BASE, "产品使用指南.docx")

ACCENT = RGBColor(0x2f, 0x6b, 0xff)
INK = RGBColor(0x1f, 0x2d, 0x4d)
SUB = RGBColor(0x5a, 0x6a, 0x88)
DANGER = RGBColor(0xc0, 0x39, 0x2b)
GREEN = RGBColor(0x1f, 0xa4, 0x63)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.5

def set_cn(run):
    run.font.name = "Microsoft YaHei"
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def shade_para(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)

def callout(label, text, kind="note"):
    colors = {
        "note": ("F5F8FF", ACCENT, "说明"),
        "warn": ("FFF5F5", DANGER, "注意"),
        "tip":  ("F3FBF6", GREEN, "提示"),
        "shot": ("FBFCFF", ACCENT, "截图"),
    }
    fill, color, tag = colors[kind]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade_cell(cell, fill)
    # thin left border accent
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0"); left.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    borders.append(left)
    tcPr.append(borders)
    p = cell.paragraphs[0]
    r = p.add_run("【%s】 " % tag); r.bold = True; r.font.color.rgb = color; set_cn(r)
    r2 = p.add_run(text); set_cn(r2)
    doc.add_paragraph()
    return tbl

def H1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); set_cn(r)
    return p

def H2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); set_cn(r)
    return p

def para(text="", bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; set_cn(r)
    if color is not None: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); set_cn(r)
    return p

# ===================== COVER =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SCAN.GATE 打印机扫描共享工具")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = ACCENT; set_cn(r)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("产品使用指南 ｜ 版本 v4.0.0 ｜ 适用对象：运营 / 行政 / 普通办公用户")
r.font.size = Pt(11); r.font.color.rgb = SUB; set_cn(r)
tagp = doc.add_paragraph()
tagp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tagp.add_run("从入门到进阶 · 图文操作手册")
r.font.size = Pt(10.5); r.font.color.rgb = ACCENT; set_cn(r)
para()
para("本文档按「从入门到进阶」顺序编排，每步均含 操作目的 · 操作方法 · 预期结果，并在关键处给出注意事项与常见问题提示。")

# ===================== 一 =====================
H1("一、产品简介与功能概述")
para("SCAN.GATE 是一款连接公区打印机网络扫描共享目录的桌面工具，帮助你在电脑上浏览、预览、上传、下载、删除扫描出来的 PDF 文件，无需手动映射网络驱动器。")

H2("核心功能")
func_rows = [
    ("服务器连接", "输入带用户名/密码的 SMB 共享路径（如 \\\\192.168.4.82\\share\\PDF），一键连接"),
    ("文件浏览", "列表展示扫描件名称、大小、修改时间，支持刷新"),
    ("文件预览", "选中 PDF 即时生成缩略预览，多页文档可翻页"),
    ("上传 / 下载", "把本地文件传入共享目录，或把扫描件保存到本机"),
    ("删除", "清理过期扫描件（带二次确认，防误删）"),
    ("多服务器管理", "保存多组服务器配置，一键切换，支持增 / 改 / 删"),
    ("会话溯源日志", "每次「连接→断开」自动在共享目录写一份中文操作日志"),
    ("窗口自由操控", "无边框毛玻璃界面，支持标题栏拖动、八向缩放、最大化"),
]
t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].paragraphs[0].add_run("功能").bold = True
hdr[1].paragraphs[0].add_run("说明").bold = True
for name, desc in func_rows:
    c = t.add_row().cells
    c[0].paragraphs[0].add_run(name).bold = True
    c[1].paragraphs[0].add_run(desc)
    for cell in c:
        for p in cell.paragraphs:
            for rr in p.runs: set_cn(rr)
doc.add_paragraph()
para("图 1 · 主界面布局：顶部标题栏（品牌 + 连接状态 + 窗口按钮）、左侧「连接设置」、中间「文件列表」、右侧「预览」、窗口四周为缩放热区。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "main-ui.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===================== 二 =====================
H1("二、环境要求与启动")
bullet("操作系统：Windows 10 / Windows 11（系统已自带 Edge WebView2 运行环境，无需额外安装）。")
bullet("网络：与目标打印机/共享服务器在同一内网，且拥有该共享目录的访问账号。")
bullet("启动方式：双击 打印机扫描工具_v4.exe 即可。程序为单实例——若已在运行，再次双击只会提示而不会打开第二个窗口。")
callout("", "程序不依赖外部浏览器，但首次启动需系统 WebView2 组件。若公司电脑禁用了系统组件导致白屏，请联系 IT 确认 WebView2 已启用。", "warn")
callout("", "建议将 打印机扫描工具_v4.exe 固定到任务栏或发送到桌面快捷方式，便于日常使用。", "tip")

# ===================== 三 =====================
H1("三、入门操作")

H2("步骤 1 · 认识主界面并连接默认服务器")
para("操作目的：建立与打印机扫描共享目录的连接，是后续所有操作的前提。", bold=True)
para("操作方法：")
bullet("程序启动后，左侧「连接设置」面板已预填默认服务器（192.168.4.82 / share / PDF / 账号 share）。")
bullet("确认信息无误，点击蓝色按钮「连接共享」。")
para("预期结果：顶部状态栏显示「已连接」，连接状态指示灯由红色变为绿色；中间「文件列表」开始加载并显示扫描件。")
callout("", "连接成功后，截取主界面，重点展示左上角连接指示灯变绿、中间文件列表出现文件行。", "shot")
para("图 3 · 步骤 1 示意：连接成功后指示灯变绿（①）、文件列表出现扫描件（②）。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step1.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "若提示连接失败，请检查：① 电脑与打印机是否在同一网段；② 共享地址/共享名/子目录/账号密码是否填写正确；③ 账号是否被服务器锁定。详见文末「故障排除」。", "warn")

H2("步骤 2 · 浏览与刷新文件列表")
para("操作目的：查看当前共享目录中有哪些扫描件。", bold=True)
para("操作方法：文件列表已自动加载；如需手动刷新，点击列表上方「刷新」按钮。")
para("预期结果：列表按「名称 / 大小 / 修改时间」三列展示所有文件；新扫描的文件在刷新后出现。")
callout("", "截取文件列表面板，展示若干 PDF 行及表头三列。", "shot")
para("图 4 · 步骤 2 示意：文件列表展示名称/大小/修改时间三列及若干 PDF 行。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step2.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "文件名带日期（如 扫描件_20260709_001.pdf）便于识别扫描时间。", "tip")

H2("步骤 3 · 预览扫描件内容")
para("操作目的：在下载/删除前先确认内容，避免误操作。", bold=True)
para("操作方法：在文件列表中单击任意文件行。")
para("预期结果：右侧「预览」面板显示该 PDF 首页图像；若为多页文档，面板底部出现「‹ 上一页 / 下一页 ›」翻页控件。")
callout("", "选中一个文件后，截取右侧预览面板显示文档图像、底部翻页控件的画面。", "shot")
para("图 5 · 步骤 3 示意：选中文件后右侧预览面板显示文档图像及翻页控件。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step3.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "文件夹（目录）不支持预览，选中后会提示「文件夹不支持预览」。", "warn")

H2("步骤 4 · 下载文件到本机")
para("操作目的：把需要的扫描件保存到自己的电脑。", bold=True)
para("操作方法：")
bullet("在文件列表中单击选中目标文件（确保不是文件夹）。")
bullet("点击「下载」按钮。")
bullet("在弹出的系统「选择保存位置」对话框中指定路径并确认。")
para("预期结果：文件保存到指定位置，进度条短暂显示后完成；状态栏提示成功。")
callout("", "截取点击「下载」后弹出的系统「另存为」对话框。", "shot")
para("图 6 · 步骤 4 示意：系统「另存为」对话框（文件名已自动填充 PDF 名称）。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step4.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "未连接共享或尚未选中文件时点击「下载」，会弹出提示「请先连接共享 / 请先选择要下载的文件」，操作不会执行。", "warn")

# ===================== 四 =====================
H1("四、进阶操作")

H2("步骤 5 · 上传文件到共享目录")
para("操作目的：将本地文件（如已签字的回执）回传到扫描共享目录，供他人取用。", bold=True)
para("操作方法：点击「上传」按钮，在文件选择对话框中选中要上传的文件（可多选），确认后等待上传完成。")
para("预期结果：上传成功后，刷新文件列表可见新文件；会话日志会记录本次上传。")
callout("", "截取点击「上传」后弹出的系统文件选择对话框（可多选状态）。", "shot")
para("图 7 · 步骤 5 示意：系统文件选择对话框（可多选状态，已勾选两份文件）。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step5.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

H2("步骤 6 · 删除过期扫描件")
para("操作目的：清理共享目录中不再需要的文件，释放空间。", bold=True)
para("操作方法：")
bullet("单击选中要删除的文件。")
bullet("点击「删除」按钮。")
bullet("在确认框中点击「确定」（此操作不可恢复）。")
para("预期结果：文件从列表中移除，并从共享目录删除；会话日志记录删除操作。")
callout("", "截取点击「删除」后弹出的确认对话框（含文件名与「不可恢复」警示）。", "shot")
para("图 8 · 步骤 6 示意：删除确认弹窗（含文件名与「不可恢复」警示）。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step6.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "删除前务必通过「预览」确认内容，避免误删他人文件。已断开连接时无法删除：指示灯为红色（未连接）时点「删除」会被拦截并提示「请先连接共享后再删除」。删除动作会写入溯源日志，操作人可追溯，请谨慎使用。", "warn")

H2("步骤 7 · 管理多组服务器配置")
para("操作目的：在多个打印机/共享位置之间快速切换，免去反复手填。", bold=True)
para("操作方法：")
bullet("点击左侧面板标题旁的「管理」按钮，打开「服务器管理」弹窗。")
bullet("在弹窗中可：连接（切换为当前使用项）、编辑（修改名称/地址/账号等）、删除（移除配置，至少保留一项）、+ 添加服务器（新建一组配置并保存）。")
bullet("完成后点击弹窗空白处或右上角 × 关闭。")
para("预期结果：配置被保存并持久化；「当前」徽标标记正在使用的服务器。")
callout("", "截取「服务器管理」弹窗，展示含「当前」徽标的配置列表与「+ 添加服务器」按钮。", "shot")
doc.add_paragraph()
para("图 2 · 服务器管理弹窗：列出已保存配置，每项含「连接 / 编辑 / 删除」，底部可新增。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "server-manage.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "把常用的公区打印机设为「当前」配置，下次启动即自动预填，直接点「连接共享」即可。", "tip")

H2("步骤 8 · 窗口移动、缩放与最大化")
para("操作目的：根据屏幕空间自由调整窗口大小与位置。", bold=True)
para("操作方法：")
bullet("移动：在顶部标题栏空白处（品牌名区域）按下并拖动。")
bullet("八向缩放：将鼠标移到窗口上 / 下 / 左 / 右 / 四角边缘，光标变为双向箭头后拖动，即可从对应方向改变大小（最小 760×480）。")
bullet("最大化 / 还原：点击右上角最大化按钮，或双击标题栏。")
bullet("最小化 / 关闭：点击右上角对应按钮。")
para("预期结果：窗口随拖动实时变化；最大化后缩放热区自动禁用，避免误触。")
callout("", "截取鼠标悬停在窗口右下角时显示斜向缩放箭头的画面。", "shot")
para("图 9 · 步骤 8 示意：窗口右下角悬停时显示斜向缩放箭头光标。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step8.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
callout("", "拖动请使用标题栏区域，不要在文件列表/按钮上拖，否则不会移动窗口。", "warn")

H2("步骤 9 · 查看「关于」与联系作者")
para("操作目的：了解版本信息，或在需要时联系程序作者。", bold=True)
para("操作方法：点击左侧「关于程序」，弹窗显示名称、版本、作者、版权；点击作者名，选择「公司内 / 公司外」身份后跳转对应飞书邀请链接。")
para("预期结果：弹出「关于 SCAN.GATE」窗口；选择身份后浏览器打开联系入口。")
callout("", "截取「关于」弹窗及随后出现的「公司内 / 公司外」选择弹窗。", "shot")
para("图 10 · 步骤 9 示意：「关于 SCAN.GATE」信息弹窗。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step9a.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("图 11 · 步骤 9 示意：身份选择弹窗（公司内 / 公司外）。示意图，以实际程序为准。", color=SUB, size=9.5)
doc.add_picture(os.path.join(IMG, "fg_step9b.png"), width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===================== 五 =====================
H1("五、会话溯源日志")
bullet("日志位置：自动写入当前所连服务器的共享根目录下 share\\log 文件夹，即 \\\\<服务器IP>\\share\\log\\。")
bullet("文件命名：log_YYYYMMDD_HHMMSS.log（以本次连接开始时间命名，精确到秒，避免重名）。")
bullet("记录内容：一次「连接 → 断开」（或关闭窗口）会话内的全部操作，含操作时间、操作人、前后状态对比、结果。")
bullet("溯源价值：操作人默认取本机 Windows 登录账号；如需显示友好姓名，可在配置文件 ~/.printer_scan_config.json 中加入 \"operator\":\"你的姓名\"。")
callout("", "日志仅在成功连接后产生操作、再断开时生成。仅浏览不操作也会生成一条含连接/断开信息的会话记录。", "tip")

# ===================== 六 =====================
H1("六、注意事项汇总")
for i, t in enumerate([
    "先连接，后操作：上传 / 下载 / 删除都要求处于「已连接」状态；断开后按钮会被拦截并提示。",
    "删除需谨慎：删除有二次确认且写入日志，操作可追溯，误删不可恢复。",
    "单实例限制：程序只能同时运行一个，重复启动会被提示。",
    "网络依赖：所有文件操作依赖内网共享可达，离网或 VPN 断开会导致连接失败。",
    "配置持久化：服务器配置保存在本机 ~/.printer_scan_config.json，重装系统前请备份。",
], 1):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(t); set_cn(r)

# ===================== 七 =====================
H1("七、故障排除（FAQ）")
faq = [
    ("双击后无窗口 / 白屏", "系统 WebView2 组件缺失或被禁用", "确认 Win10/11 已启用 Edge WebView2；联系 IT 修复系统组件"),
    ("提示「连接失败」", "地址/账号错误、不在同网段、服务器不可达", "核对 IP、共享名、子目录、用户名、密码；用 ping <IP> 测试连通性"),
    ("文件列表为空", "已连接但目录无文件，或子目录填错", "检查「子目录」是否应为空或具体文件夹名（如 PDF）"),
    ("点「下载/删除」没反应并弹提示", "未连接或没选中文件", "先「连接共享」，再在列表单击选中具体文件"),
    ("指示灯是红色仍想操作", "当前处于「未连接」状态", "重新点击「连接共享」；断开后列表已清空属正常"),
    ("重建后 exe 图标没变", "Windows 资源管理器图标缓存未刷新", "桌面按 F5 刷新，或重启「资源管理器」，或执行 ie4uinit.exe -show 清缓存"),
    ("提示「程序已在运行」", "已有一个实例", "切到已打开的窗口，或结束后台进程后再启动"),
    ("日志没生成", "未成功连接，或共享 log 目录无写权限", "确认先连接并产生操作；检查对 \\<IP>\\share 是否有写入权限"),
    ("预览不显示", "文件非 PDF / 文件损坏 / 正在生成", "刷新重试；非 PDF 文件不支持预览"),
]
ft = doc.add_table(rows=1, cols=3); ft.style = "Light Grid Accent 1"
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
hc = ft.rows[0].cells
for i, h in enumerate(["现象", "可能原因", "解决办法"]):
    hc[i].paragraphs[0].add_run(h).bold = True
for a, b, c in faq:
    row = ft.add_row().cells
    row[0].paragraphs[0].add_run(a)
    row[1].paragraphs[0].add_run(b)
    row[2].paragraphs[0].add_run(c)
    for cell in row:
        for p in cell.paragraphs:
            for rr in p.runs: set_cn(rr)
doc.add_paragraph()
callout("", "若以上方法无法解决，请将现象截图及 \\<IP>\\share\\log 下对应日志文件一并反馈给作者。", "note")

# footer
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("— SCAN.GATE 打印机扫描共享工具 · 产品使用指南 v4.0.0 —")
r.font.size = Pt(9); r.font.color.rgb = SUB; set_cn(r)

doc.save(OUT)
print("SAVED:", OUT, os.path.getsize(OUT), "bytes")
